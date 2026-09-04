"""Resume parsing: markdown/text first, PDF via pdfplumber when available.

Produces a ResumeDocument where bullets are split on markdown bullets,
dash lines, or sentence boundaries. Skills default to profile skills —
the base resume is the source of truth for facts, never the LLM.
"""

from __future__ import annotations

import re
from pathlib import Path

from hermes.config import Profile
from hermes.models import Bullet, ResumeDocument

_BULLET_RE = re.compile(r"^\s*(?:[-*•]|\d+\.)\s+(.+)$", re.MULTILINE)
_HEADER_RE = re.compile(r"^(#{1,4})\s+(.+)$", re.MULTILINE)


def _extract_text_pdf(path: Path) -> str:
    import pdfplumber

    pages = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            pages.append(page.extract_text() or "")
    return "\n".join(pages)


def _extract_text_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "") if para.style else ""
        # Render list paragraphs and headings as markdown bullets/headers
        # so downstream bullet-splitting works on DOCX resumes too.
        if "Heading 1" in style:
            parts.append(f"# {text}")
        elif "Heading 2" in style or "Heading 3" in style:
            parts.append(f"## {text}")
        elif "List" in style:
            parts.append(f"- {text}")
        else:
            parts.append(text)
    return "\n\n".join(parts)


def load_resume_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text = _extract_text_pdf(path)
    elif suffix == ".docx":
        text = _extract_text_docx(path)
    else:
        text = path.read_text(encoding="utf-8", errors="replace")
    return text.strip()


_STORY_SECTIONS = (
    "experience", "projects", "research", "publications",
    "research & publications", "work", "employment",
)


def _split_bullets_section_aware(text: str) -> list[str]:
    """Extract bullets only from experience-like markdown sections."""
    section = ""       # active ## section (what we care about)
    for line in text.splitlines():
        stripped = line.strip()
        header = _HEADER_RE.match(stripped)
        if header:
            level, name = len(header.group(1)), header.group(2).strip().lower()
            if level <= 2:  # ## section switch — ### subheads don't reset it
                section = name
            continue
        bullet = _BULLET_RE.match(stripped)
        if not bullet:
            continue
        if section and not any(s in section for s in _STORY_SECTIONS):
            continue  # bullet under Skills/Education/etc — keep out of RAG
        content = bullet.group(1).strip()
        # Skip bare skill-category lines ("Category: a, b, c, d")
        if re.match(r"^[A-Z][\w&/# -]{2,40}:\s", content) and content.count(",") >= 3:
            continue
        yield content


def _split_bullets(text: str) -> list[str]:
    found = list(_split_bullets_section_aware(text))
    if found:
        return found
    # Fallback for unstructured resumes: bullets anywhere.
    found = [m.group(1).strip() for m in _BULLET_RE.finditer(text)]
    if found:
        return found
    # Final fallback: sentence split.
    return [
        s.strip()
        for s in re.split(r"(?<=[.!?])\s+", text)
        if len(s.strip()) > 30
    ]


def _guess_seniority(text: str, profile: Profile) -> str:
    lowered = text.lower()
    if profile.target.seniority:
        return profile.target.seniority
    if "staff" in lowered or "principal" in lowered:
        return "staff"
    if "senior" in lowered:
        return "senior"
    if "junior" in lowered or "intern" in lowered:
        return "junior"
    return "mid"


def parse_resume(path: Path, profile: Profile) -> ResumeDocument:
    text = load_resume_text(path)
    return parse_resume_text(text, profile)


def parse_resume_text(text: str, profile: Profile) -> ResumeDocument:
    """Parse raw resume text (from file or web paste) into a ResumeDocument."""
    bullets_raw = _split_bullets(text)
    profile_skills = profile.all_skills

    bullets = [
        Bullet(
            id=f"bullet-{i:03d}",
            text=bullet,
            skills=[
                skill
                for skill in profile_skills
                if skill.lower() in bullet.lower()
            ],
        )
        for i, bullet in enumerate(bullets_raw)
    ]

    return ResumeDocument(
        name=profile.identity.name,
        email=profile.identity.email,
        phone=profile.identity.phone,
        summary=_extract_summary(text),
        skills=profile_skills,
        bullets=bullets,
        seniority=_guess_seniority(text, profile),
        years_experience=profile.target.years_experience,
        raw_text=text,
    )


def _extract_summary(text: str) -> str:
    """Grab the first non-heading paragraph as a summary proxy."""
    lines = text.splitlines()
    paragraph: list[str] = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            if paragraph:
                break
            continue
        if stripped.startswith("#"):
            if paragraph:
                break
            continue
        if stripped.startswith(("-", "*", "•")):
            break
        paragraph.append(stripped)
        if sum(len(p) for p in paragraph) > 600:
            break
    return " ".join(paragraph)[:600]
